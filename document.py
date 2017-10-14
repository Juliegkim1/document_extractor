import docx
import string


class Document:
    """ structure the text from a single document
     Document includes separate methods to extract and structure
     sections, tables, and properties from a docx file

     NOTE: only compatible with docx, not earlier version like doc
     """

    def __init__(self, path, doc_text=True, sections=True, table_text=True, doc_properties=True):
        """
        :param path: directory path to document
        :param doc_text: if True, extract document text
        :param sections: if True, structure document text into sections
        :param doc_properties: if True, extract document properties (e.g. author)
        :param table_text: if True, extract the text from document tables
        """

        # use docx to read document xml
        self.path = path
        doc = docx.Document(self.path)
        self.sections = None
        self.table_text = None

        if doc_text:
            self.paragraphs = doc.paragraphs
            self.text = ' '.join([para.text.strip() for para in self.paragraphs if para.text.strip() != ''])

        if sections:
            self.set_sections(doc)

        if table_text:
            self.set_table_text(doc)

        if doc_properties:
            self.author = doc.core_properties.author
            self.last_modified_by = doc.core_properties.last_modified_by
            self.created = doc.core_properties.created
            self.last_printed = doc.core_properties.last_printed
            self.revision = doc.core_properties.revision
            self.num_tables = len(doc.tables)

    def __repr__(self):
        return "<DOCUMENT: {}>".format(self.path)

    def set_table_text(self, doc):
        """ set the table_text attribute with text from the documents tables

        :param doc: docx document
        """

        table_text = []
        for table in doc.tables:
            for cell in table._cells:
               table_text.append(cell.text.strip())
        self.table_text = ' '.join(table_text)

    def set_sections(self
                     , doc
                     , use_headings=True
                     , use_capitalization=True
                     , use_bold=True
                     , use_underline=True
                     , use_bold_until_colon=True
                     , use_capital_letter_list=True
                     , use_roman_numeral_list=True
                     , ignore_bullets=True):

        """ set sections using current heuristics to detect section headers

        :param use_headings: uses a header formatting (e.g. table of contents)
        :param use_capitalization: capitalization of every letter often indicates section header
        :param use_bold: all words in a sentence are bold
        :param use_underline: all words in a sentence are underlined
        :param use_bold_until_colon: all words in a sentence are bold until a colon (e.g. SECTION: ...)
        :param use_capital_letter_list: sentence starts with capital letter (e.g. A.)
        :param use_roman_numeral_list: sentence starts with a roman numeral (e.g. II.)
        :param ignore_bullets: ignore list of bullet points
        """

        # add structure for document sections
        sections = {}
        section_text = []

        # set a name for the first section in case
        # none of the section criteria are met in the document
        section_name = ''

        for p in doc.paragraphs:

            # docx stores many paragraph which are not informative
            # e.g. line breaks
            if self.paragraph_doesnt_have_text(p):
                continue

            if not self.is_section_header(p
                                          , use_headings
                                          , use_capitalization
                                          , use_bold
                                          , use_underline
                                          , use_bold_until_colon
                                          , use_capital_letter_list
                                          , use_roman_numeral_list
                                          , ignore_bullets):
                section_text.append(p.text)
            else:
                text = ' '.join(section_text).strip()

                section_has_text = text != ''
                if section_has_text:
                    # populate sections dict with the completed section
                    sections[section_name] = text
                    # reset name and text for next section
                    section_name = p.text.upper().strip()
                    section_text = []
                else:
                    # section name is spread across multiple lines
                    section_name = ' '.join([section_name, p.text.upper().strip()])

        # add the text from the final section to sections dict
        text = ' '.join(section_text).strip()
        section_has_text = text != ''
        if section_has_text:
            sections[section_name] = text

        self.sections = sections

    @staticmethod
    def paragraph_doesnt_have_text(p, alpha_only=True):
        """ ignore paragraphs that do not contain any text

        :param p: paragraph
        :param alpha only: if True, only keep paragraph that have at least one letter (e.g. ignore phone #)
        :return: bool (True) if the paragraph is empty
        """

        empty_string = p.text.strip() == ''

        # ignores phone numbers and string of non-text characters (e.g. _____)
        has_letters = True
        if alpha_only:
            has_letters = any(char.isalpha() for char in p.text)

        if empty_string or not has_letters:
            return True
        return False

    @staticmethod
    def is_section_header(p
                          , use_headings=True
                          , use_capitalization=True
                          , use_bold=True
                          , use_underline=True
                          , use_bold_until_colon=True
                          , use_capital_letter_list=True
                          , use_roman_numeral_list=True
                          , ignore_bullets=True):

        """ determine if a paragraph is a section header

        :param p: paragraph
        :return section_header: boolean if paragraph is a section header

        NOTE: there is not an exact method to determine a section header
        due to inconsistencies in the way documents are created.
        This function uses heuristics (e.g. all CAPS) to determine sections

        """

        section_header = False

        # uses Header formatting (common for table of contents section headers)
        if use_headings:
            if 'HEADING' in p.style.name.upper():
                section_header = True

        # capitalization of every letter often indicates section header
        if use_capitalization:
            if p.text.isupper():
                section_header = True

        # check for bold, underline, and colon conditions
        bold_runs = []
        underline_runs = []
        colon_runs = []
        colon_continue = True
        for run in p.runs:
            # ignore runs (style) for blank space at end of sentence
            if run.text.strip() == '':
                continue
            if use_bold:
                bold_runs.append(run.bold)
            if use_underline:
                underline_runs.append(run.underline)
            if use_bold_until_colon:
                # keep adding runs until a colon is found
                # for sections that have bold text until colon
                # (e.g. SECTION: ...)
                if colon_continue:
                    colon_runs.append(run.bold)
                    if ':' in run.text:
                        colon_continue = False
        bold_cond = all(bold_runs) and bold_runs != list()
        underline_cond = all(underline_runs) and underline_runs != list()
        colon_cond = all(colon_runs) and colon_runs != list()
        if bold_cond or underline_cond or colon_cond:
            section_header = True

        # find list items that start with a capital letter
        # e.g.  A. section text
        #       B. section text
        #       C. section text
        if use_capital_letter_list:
            upper_case_letters = [''.join([char, '. ']) for char in string.ascii_uppercase]
            upper_case_letter_list = (p.text.strip()[0:3] in upper_case_letters)
            if upper_case_letter_list:
                section_header = True

        # find list items that start with a roman_numeral
        # e.g.  I.   section text
        #       II.  section text
        #       III. section text
        if use_roman_numeral_list:
            # only look at the characters necessary to identify roman numeral
            one_letter_numeral = (p.text.strip()[0:3] in ['I. ', 'V. ', 'X. '])
            two_letter_numeral = (p.text.strip()[0:4] in ['II. ', 'IV. ', 'VI. ', 'IX. ',  'XI. '])
            three_letter_numeral = (p.text.strip()[0:5] in [ 'III. ', 'VII. ', 'XII. '])
            four_letter_numeral = (p.text.strip()[0:6] in ['VIII.'])
            roman_numeral_start = (one_letter_numeral or two_letter_numeral
                                   or three_letter_numeral or four_letter_numeral)
        if roman_numeral_start:
            section_header = True

        # bullets points are often converted into the letter 'O' followed by a space
        if ignore_bullets:
            if p.text.strip()[0:2] in ['o ', 'O ']:
                section_header = False

        return section_header

